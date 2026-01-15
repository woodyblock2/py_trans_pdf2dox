from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


class DocxWriter:
    def __init__(self) -> None:
        self.document = Document()

    def _set_font(self, run) -> None:
        for font_name in ["SimSun", "Microsoft YaHei"]:
            try:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                break
            except Exception:  # noqa: BLE001
                continue
        run.font.size = Pt(10)

    def add_table(self, table_data: dict[str, Any]) -> None:
        rows = table_data["rows"]
        cols = table_data["cols"]
        if rows <= 0 or cols <= 0:
            return
        table = self.document.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        cell_map = {(cell["row"], cell["col"]): cell for cell in table_data["cells"]}

        for cell in table_data["cells"]:
            row = cell["row"]
            col = cell["col"]
            rowspan = cell.get("rowspan", 1)
            colspan = cell.get("colspan", 1)
            target = table.cell(row, col)
            if rowspan > 1 or colspan > 1:
                target = target.merge(table.cell(row + rowspan - 1, col + colspan - 1))
            target.text = cell.get("text", "")
            if target.paragraphs and target.paragraphs[0].runs:
                self._set_font(target.paragraphs[0].runs[0])

    def add_page_break(self) -> None:
        self.document.add_page_break()

    def save(self, path: Path) -> None:
        self.document.save(path)
