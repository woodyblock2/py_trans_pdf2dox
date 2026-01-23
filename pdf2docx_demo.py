#!/usr/bin/env python3
import argparse
import inspect
import logging
import os
from pathlib import Path

import fitz
import importlib
import numpy as np
from bs4 import BeautifulSoup
from docx import Document


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="PDF→DOCX demo using PaddleOCR PP-OCRv5 pipeline (PPStructure)."
    )
    parser.add_argument("pdf", help="Path to input PDF")
    parser.add_argument(
        "-o",
        "--output",
        help="Path to output DOCX (default: same name as PDF)",
    )
    parser.add_argument(
        "--dpi",
        type=int,
        default=300,
        help="Render DPI for PDF pages (default: 300)",
    )
    parser.add_argument(
        "--offline",
        action="store_true",
        help="Offline mode: require local models, no downloads",
    )
    parser.add_argument("--lang", default="ch", help="OCR language (default: ch)")
    parser.add_argument("--det-model-dir", help="Detection model directory")
    parser.add_argument("--rec-model-dir", help="Recognition model directory")
    parser.add_argument("--cls-model-dir", help="Classifier model directory")
    parser.add_argument("--table-model-dir", help="Table model directory")
    parser.add_argument("--layout-model-dir", help="Layout model directory")
    parser.add_argument(
        "--structure-types",
        default="table,text",
        help="Comma-separated structure types for PPStructure (default: table,text)",
    )
    parser.add_argument(
        "--show-log",
        action="store_true",
        help="Show PaddleOCR pipeline logs",
    )
    return parser.parse_args()


def ensure_offline_models(args: argparse.Namespace) -> None:
    missing = []
    for flag, value in (
        ("--det-model-dir", args.det_model_dir),
        ("--rec-model-dir", args.rec_model_dir),
        ("--cls-model-dir", args.cls_model_dir),
        ("--table-model-dir", args.table_model_dir),
        ("--layout-model-dir", args.layout_model_dir),
    ):
        if not value:
            missing.append(flag)
        else:
            if not Path(value).exists():
                raise FileNotFoundError(
                    f"Offline mode requires local model dir {value}, but it does not exist."
                )
    if missing:
        raise ValueError(
            "Offline mode requires explicit local model directories. "
            f"Missing: {', '.join(missing)}"
        )


def resolve_ppstructure_class():
    try:
        module = importlib.import_module("paddleocr.ppstructure")
        return module.PPStructure
    except ModuleNotFoundError:
        module = importlib.import_module("paddleocr")
        if not hasattr(module, "PPStructure"):
            raise ModuleNotFoundError(
                "PPStructure not found in paddleocr. "
                "Please install PaddleOCR 2.7+ and ensure ppstructure module is available."
            )
        return module.PPStructure


def build_ppstructure(args: argparse.Namespace):
    kwargs = {
        "show_log": args.show_log,
        "lang": args.lang,
    }
    if args.det_model_dir:
        kwargs["det_model_dir"] = args.det_model_dir
    if args.rec_model_dir:
        kwargs["rec_model_dir"] = args.rec_model_dir
    if args.cls_model_dir:
        kwargs["cls_model_dir"] = args.cls_model_dir
    if args.table_model_dir:
        kwargs["table_model_dir"] = args.table_model_dir
    if args.layout_model_dir:
        kwargs["layout_model_dir"] = args.layout_model_dir

    structure_types = [s.strip() for s in args.structure_types.split(",") if s.strip()]
    if structure_types:
        kwargs["structure_type"] = structure_types

    PPStructure = resolve_ppstructure_class()
    sig = inspect.signature(PPStructure.__init__)
    if "ocr_version" in sig.parameters:
        kwargs["ocr_version"] = "PP-OCRv5"
    if "use_angle_cls" in sig.parameters:
        kwargs["use_angle_cls"] = True
    if "use_textline_orientation" in sig.parameters:
        kwargs["use_textline_orientation"] = True

    logging.info("Initializing PPStructure with PP-OCRv5 pipeline")
    return PPStructure(**kwargs)


def render_pdf_pages(pdf_path: Path, dpi: int):
    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)
    with fitz.open(pdf_path) as doc:
        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=matrix)
            image = np.frombuffer(pix.samples, dtype=np.uint8)
            image = image.reshape(pix.height, pix.width, pix.n)
            if pix.n == 4:
                image = image[:, :, :3]
            image = image[:, :, ::-1]
            yield page_index + 1, doc.page_count, image


def bbox_key(element):
    bbox = element.get("bbox") or element.get("box")
    if not bbox:
        return (0, 0)
    if len(bbox) == 4:
        x1, y1, x2, y2 = bbox
    else:
        xs = [p[0] for p in bbox]
        ys = [p[1] for p in bbox]
        x1, y1 = min(xs), min(ys)
        x2, y2 = max(xs), max(ys)
    return (y1, x1)


def extract_text_lines(res):
    lines = []
    if isinstance(res, str):
        return [res]
    if isinstance(res, list):
        for item in res:
            if isinstance(item, dict):
                text = item.get("text")
                if text:
                    lines.append(text)
            elif isinstance(item, (list, tuple)):
                if len(item) == 2:
                    text_part = item[1]
                    if isinstance(text_part, (list, tuple)) and text_part:
                        lines.append(str(text_part[0]))
                    elif isinstance(text_part, str):
                        lines.append(text_part)
                elif len(item) == 1 and isinstance(item[0], str):
                    lines.append(item[0])
            elif isinstance(item, str):
                lines.append(item)
    return lines


def append_table_from_html(document: Document, html: str) -> None:
    soup = BeautifulSoup(html, "html.parser")
    table_tag = soup.find("table")
    if not table_tag:
        document.add_paragraph(html)
        return

    rows = table_tag.find_all("tr")
    if not rows:
        return

    max_cols = 0
    for row in rows:
        cells = row.find_all(["td", "th"])
        count = 0
        for cell in cells:
            colspan = int(cell.get("colspan", 1))
            count += colspan
        max_cols = max(max_cols, count)

    docx_table = document.add_table(rows=len(rows), cols=max_cols)

    for row_idx, row in enumerate(rows):
        col_idx = 0
        cells = row.find_all(["td", "th"])
        for cell in cells:
            while col_idx < max_cols and docx_table.cell(row_idx, col_idx).text:
                col_idx += 1
            cell_text = cell.get_text(strip=True)
            colspan = int(cell.get("colspan", 1))
            rowspan = int(cell.get("rowspan", 1))

            target_cell = docx_table.cell(row_idx, col_idx)
            target_cell.text = cell_text

            if colspan > 1:
                merge_to = docx_table.cell(row_idx, col_idx + colspan - 1)
                target_cell = target_cell.merge(merge_to)
            if rowspan > 1:
                merge_to = docx_table.cell(row_idx + rowspan - 1, col_idx)
                target_cell.merge(merge_to)
            col_idx += colspan


def append_structure_to_docx(document: Document, structure_result: list) -> None:
    for element in sorted(structure_result, key=bbox_key):
        element_type = element.get("type")
        if element_type == "table":
            html = element.get("res", {}).get("html") if isinstance(element.get("res"), dict) else element.get("html")
            if html:
                append_table_from_html(document, html)
            else:
                document.add_paragraph("[table]")
        else:
            res = element.get("res", [])
            lines = extract_text_lines(res)
            if lines:
                for line in lines:
                    document.add_paragraph(line)
            else:
                text = element.get("text")
                if text:
                    document.add_paragraph(text)


def main() -> None:
    args = parse_args()
    logging.basicConfig(level=logging.INFO, format="%(message)s")

    pdf_path = Path(args.pdf)
    if not pdf_path.exists():
        raise FileNotFoundError(f"Input PDF not found: {pdf_path}")

    output_path = Path(args.output) if args.output else pdf_path.with_suffix(".docx")

    if args.offline:
        ensure_offline_models(args)

    ocr_engine = build_ppstructure(args)
    document = Document()

    for page_number, total_pages, image in render_pdf_pages(pdf_path, args.dpi):
        logging.info("Processing page %s", page_number)
        result = ocr_engine(image)
        append_structure_to_docx(document, result)
        if page_number < total_pages:
            document.add_page_break()

    document.save(output_path)
    logging.info("Saved DOCX to %s", output_path)


if __name__ == "__main__":
    main()
